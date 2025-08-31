import asyncio
import logging
from pathlib import Path
from typing import List, Union

import docx

from ..dataclasses import ParsedContent, FileType
from .base import BaseParser
from .cleaners import clean_text


class DocxParser(BaseParser):
    def __init__(self, **kwargs):
        self.logger = logging.getLogger(__name__)

    def get_file_extensions(self) -> List[str]:
        return FileType.extensions[FileType.DOCX]

    def get_file_type(self) -> str:
        return FileType.DOCX

    async def parse_async(self, file_path: Union[str, Path]) -> ParsedContent:
        return await asyncio.to_thread(self._parse, file_path)

    def _parse(self, file_path: Union[str, Path]) -> ParsedContent:
        file_path = self._validate_file(file_path)

        try:
            doc = docx.Document(file_path)
        except Exception as e:
            self.logger.error(f"Failed to open DOCX file: {e}")
            raise

        text_parts = []
        tables = []

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text if cell_text else "")
                if any(row_data):
                    table_data.append(row_data)
            if table_data:
                tables.append(table_data)
                text_parts.append("[TABLE]")

        current_heading = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Xác định kiểu paragraph
            style = para.style.name if para.style else ""

            is_heading = (
                style.startswith('Heading') or
                style == 'Title' or
                (len(text) < 100 and text.isupper() and text.replace(' ', '').isalpha())
            )

            if is_heading:
                if current_heading:
                    text_parts.append(f"[HEADING]{current_heading}[/HEADING]")
                current_heading = text
            else:
                if current_heading:
                    text_parts.append(f"[HEADING]{current_heading}[/HEADING]")
                    current_heading = None
                text_parts.append(text)

        # Đóng heading cuối
        if current_heading:
            text_parts.append(f"[HEADING]{current_heading}[/HEADING]")

        full_text = "\n".join(filter(None, text_parts))
        full_text = clean_text(full_text)

        metadata = {
            'file_name': file_path.name,
            'file_size': file_path.stat().st_size,
            'parser': 'docx_parser',
            'num_tables': len(tables)
        }

        return ParsedContent(
            content=full_text,
            metadata=metadata,
            file_type=self.get_file_type(),
            file_path=str(file_path),
            tables=tables
        )